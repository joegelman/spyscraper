#!/usr/bin/env python3
import argparse
import json
import os
import ssl
import time
from pathlib import Path
from urllib.error import HTTPError
from urllib.request import Request, urlopen

from docx import Document
from openpyxl import load_workbook

# -----------------------
# JSON SCHEMAS
# -----------------------

FULL_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "required": [
        "title",
        "subtitle",
        "executive_summary",
        "key_differentiators",
        "platform_stack",
        "core_capabilities",
        "product_categories",
    ],
    "properties": {
        "title": {"type": "string", "minLength": 3},
        "subtitle": {"type": "string", "minLength": 3},
        "executive_summary": {"type": "string", "minLength": 30},
        "key_differentiators": {
            "type": "array",
            "minItems": 3,
            "items": {"type": "string", "minLength": 8},
        },
        "platform_stack": {
            "type": "array",
            "minItems": 4,
            "maxItems": 4,
            "items": {
                "type": "object",
                "additionalProperties": False,
                "required": ["layer", "name", "description"],
                "properties": {
                    "layer": {"type": "string"},
                    "name": {"type": "string"},
                    "description": {"type": "string"},
                },
            },
        },
        "core_capabilities": {
            "type": "array",
            "minItems": 5,
            "items": {"type": "string"},
        },
        "product_categories": {
            "type": "array",
            "minItems": 4,
            "items": {
                "type": "object",
                "additionalProperties": False,
                "required": ["name", "url", "sections"],
                "properties": {
                    "name": {"type": "string"},
                    "url": {"type": "string"},
                    "sections": {
                        "type": "array",
                        "minItems": 2,
                        "items": {
                            "type": "object",
                            "additionalProperties": False,
                            "required": ["heading", "bullets"],
                            "properties": {
                                "heading": {"type": "string"},
                                "bullets": {
                                    "type": "array",
                                    "minItems": 2,
                                    "items": {"type": "string"},
                                },
                            },
                        },
                    },
                },
            },
        },
    },
}

FALLBACK_SCHEMA = {
    "type": "object",
    "additionalProperties": False,
    "required": ["title", "subtitle", "executive_summary", "product_categories"],
    "properties": {
        "title": {"type": "string"},
        "subtitle": {"type": "string"},
        "executive_summary": {"type": "string"},
        "product_categories": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "required": ["name", "summary"],
                "properties": {
                    "name": {"type": "string"},
                    "summary": {"type": "string"},
                },
            },
        },
    },
}


# -----------------------
# OPENAI CALL
# -----------------------


def call_openai_structured(prompt: str, model: str, schema: dict):
    ctx = (
        ssl._create_unverified_context()
        if os.getenv("OPENAI_INSECURE")
        else ssl.create_default_context()
    )

    body = {
        "model": model,
        "input": [
            {
                "role": "system",
                "content": "You must output ONLY a single valid JSON object. No prose.",
            },
            {"role": "user", "content": prompt},
        ],
        "text": {
            "format": {
                "type": "json_schema",
                "name": "competitive_offering_map",
                "schema": schema,
                "strict": True,
            }
        },
        "max_output_tokens": 1800,
    }

    req = Request(
        "https://api.openai.com/v1/responses",
        data=json.dumps(body).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {os.environ['OPENAI_API_KEY']}",
            "Content-Type": "application/json",
        },
    )

    raw = None
    data = None

    try:
        with urlopen(req, timeout=900, context=ctx) as resp:
            raw = resp.read().decode("utf-8")
            data = json.loads(raw)
    except HTTPError as e:
        print(e.read().decode("utf-8", errors="replace"))
        raise

    # Detect truncation
    if data.get("status") == "incomplete":
        raise RuntimeError("INCOMPLETE_OUTPUT")

    # Extract text payload
    for item in data.get("output", []):
        if item.get("type") == "message":
            for c in item.get("content", []):
                if c.get("type") == "output_text":
                    return json.loads(c["text"])

    raise RuntimeError("No JSON content found")


# -----------------------
# DOCX BUILDER
# -----------------------


def build_docx(report: dict, out_path: str):
    doc = Document()
    doc.add_heading(report["title"], 0)
    doc.add_paragraph(report["subtitle"])

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(report["executive_summary"])

    for cat in report.get("product_categories", []):
        doc.add_heading(cat["name"], level=1)
        doc.add_paragraph(cat.get("summary", ""))

    doc.save(out_path)


# -----------------------
# MAIN
# -----------------------


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--vendor", required=True)
    ap.add_argument("--domain", required=True)
    ap.add_argument("--xlsx", required=True)
    ap.add_argument("--out", required=True)
    ap.add_argument("--model", default="gpt-5-mini")
    ap.add_argument("--max-rows", type=int, default=10)
    args = ap.parse_args()

    wb = load_workbook(args.xlsx)
    ws = wb.active

    rows = []
    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True)):
        if i >= args.max_rows:
            break
        rows.append(str(row[3])[:1500])

    prompt = f"""
Using the following source material, produce a competitive offering map for {args.vendor} ({args.domain}).

SOURCE MATERIAL:
{chr(10).join(rows)}
"""

    try:
        report = call_openai_structured(prompt, args.model, FULL_SCHEMA)
    except RuntimeError:
        print("Full schema failed, retrying with fallback schema...")
        report = call_openai_structured(prompt, args.model, FALLBACK_SCHEMA)

    build_docx(report, args.out)
    print(f"Wrote {args.out}")


if __name__ == "__main__":
    main()
